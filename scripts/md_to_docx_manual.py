# -*- coding: utf-8 -*-
"""
将 standard-manual 下的 Markdown 转为 Word（.docx），并嵌入相对路径图片。
依赖: pip install python-docx
用法（在仓库根目录）:
  python scripts/md_to_docx_manual.py
  python scripts/md_to_docx_manual.py docs/standard-manual/01-求职者操作手册.md
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    from docx.shared import RGBColor
except ImportError:
    print("请先安装: pip install python-docx", file=sys.stderr)
    sys.exit(1)

IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")


def strip_md_inline(s: str) -> str:
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s.strip()


def add_runs_to_paragraph(p, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        else:
            p.add_run(part)


def add_paragraph_with_bold(doc: Document, text: str) -> None:
    text = text.rstrip()
    if not text:
        return
    p = doc.add_paragraph()
    add_runs_to_paragraph(p, text)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            table.rows[ri].cells[ci].text = strip_md_inline(cell)


def parse_cell_line(line: str) -> list[str] | None:
    line = line.strip()
    if not line.startswith("|") or not line.endswith("|"):
        return None
    inner = line[1:-1]
    return [c.strip() for c in inner.split("|")]


def md_to_docx(md_path: Path, out_dir: Path) -> Path:
    base = md_path.parent
    md_text = md_path.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    doc = Document()
    sect = doc.sections[0]
    sect.page_height = sect.page_height
    sect.page_width = sect.page_width

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip() == "---":
            if doc.paragraphs:
                doc.add_paragraph()
            i += 1
            continue

        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            title = strip_md_inline(m.group(2))
            doc.add_heading(title, level=min(level, 3))
            i += 1
            continue

        img_m = IMG_RE.search(line)
        if img_m and line.strip().startswith("!["):
            alt = img_m.group(1)
            rel = img_m.group(2).strip()
            img_path = (base / rel).resolve()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img_path.is_file():
                run = p.add_run()
                try:
                    run.add_picture(str(img_path), width=Inches(6.2))
                except Exception as ex:
                    p.add_run(f"[图片加载失败: {img_path} ({ex})]")
            else:
                p.add_run(f"[图片不存在: {img_path}]")
            if alt:
                cap = doc.add_paragraph(alt)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in cap.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        if line.strip().startswith("|") and i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            if next_line.startswith("|") and "---" in next_line:
                rows_raw: list[list[str]] = []
                head = parse_cell_line(lines[i])
                if head:
                    rows_raw.append(head)
                i += 2
                while i < len(lines) and lines[i].strip().startswith("|"):
                    cells = parse_cell_line(lines[i])
                    if cells:
                        rows_raw.append(cells)
                    i += 1
                if rows_raw:
                    add_table(doc, rows_raw)
                continue

        if line.strip().startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            add_runs_to_paragraph(p, line.strip()[2:])
            i += 1
            continue

        if line.strip() == "":
            i += 1
            continue

        add_paragraph_with_bold(doc, line)
        i += 1

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / (md_path.stem + ".docx")
    doc.save(str(out_path))
    return out_path


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "docs" / "standard-manual" / "word"
    if len(sys.argv) > 1:
        paths = [Path(sys.argv[1])]
    else:
        manual_dir = root / "docs" / "standard-manual"
        paths = sorted(manual_dir.glob("[0-9][0-9]-*.md"))
    for p in paths:
        if not p.is_file():
            print(f"跳过（不存在）: {p}", file=sys.stderr)
            continue
        out = md_to_docx(p.resolve(), out_dir)
        print(f"已生成: {out}")


if __name__ == "__main__":
    main()
