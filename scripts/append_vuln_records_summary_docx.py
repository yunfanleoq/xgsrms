# -*- coding: utf-8 -*-
"""在汇总漏洞 DOCX 主表末尾（补充条之前）追加 VULN-038～041，并将「共 37 条」改为「共 41 条」。依赖 python-docx。"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_PATH = Path(__file__).resolve().parent.parent / "Fw_ 招聘系统漏洞报告及测试记录" / "招聘系统-漏洞记录及修复情况记录表-汇总.docx"

COMMON_HOST = "10.28.12.15（渗透测试目标，见专项报告）"
COMMON_PATH = "（详见专项报告「受影响组件/URL」或 PoC）"
COMMON_SYS = "信工所招聘管理系统（JeecgBoot / xgszp-boot）"
COMMON_RESP = "刘彦良"
COMMON_VENDOR = "中科同为科技（北京）有限责任公司"
COMMON_BRANCH = "xgszp-boot / xgsboot-vue3 当前开发分支"


def base_rows(
    name: str,
    vuln_id: str,
    level: str,
    description: str,
    suggestion: str,
    proof: str,
    fix_detail: str,
    files: str,
    self_test: str,
    retest: str,
) -> list[tuple[str, str]]:
    return [
        ("漏洞基本信息", ""),
        ("漏洞名称", name),
        ("漏洞编号", vuln_id),
        ("影响主机", COMMON_HOST),
        ("漏洞路径", COMMON_PATH),
        ("业务系统", COMMON_SYS),
        ("漏洞等级", level),
        ("漏洞描述", description),
        ("修复建议", suggestion),
        ("漏洞证明", proof),
        ("修复与验证信息", ""),
        ("系统负责人", COMMON_RESP),
        ("研发厂商", COMMON_VENDOR),
        ("修复负责人", COMMON_RESP),
        ("修复方案说明", fix_detail),
        ("涉及修改的文件/代码", files),
        ("修复版本/补丁号", COMMON_BRANCH),
        ("修复时间", "待上线后填写"),
        ("上线时间", "待上线后填写"),
        ("修复方自测说明", self_test),
        ("修复确认和闭环", ""),
        ("复测结果", retest),
        ("漏洞发现人确认", "待填写"),
        ("安全团队\n确认", "待填写"),
        ("系统负责人确认", "待填写"),
        ("最终确认\n意见", "☐ 同意闭环  ☐ 退回重新修复"),
        ("完成时间", "待填写"),
    ]


SECTIONS = [
    (
        "信工所招聘管理系统 — 漏洞记录及修复情况记录表（第 38 / 41 条｜已修复）",
        base_rows(
            name="登录口令 AES 密钥匿名下发（getEncryptedString）",
            vuln_id="VULN-038",
            level="中危",
            description=(
                "【修复状态：已修复】"
                "GET /sys/getEncryptedString 为 anon，长期返回框架默认 AES key/iv，任意客户端可取密钥解密前端 AES 口令传输（与签名密钥加固 VULN-004 不同项）。"
                "【本次加固】① 新增 GET /sys/loginRsaPublicKey：RSA2048，私钥 PKCS#8 仅存 Redis（TTL 120s），服务端解密后即删除密钥；"
                "② Vue3 登录流程先取公钥，用 jsencrypt RSA/PKCS1 加密密码并提交 rsaKeyId；"
                "③ LoginController /login、/mLogin 优先 RSA 解密，无 rsaKeyId 时仍走 AesEncryptUtil.resolvePassword 兼容旧客户端；"
                "④ getEncryptedString 标注废弃；⑤ Shiro 将 loginRsaPublicKey 设为 anon。\n"
                "发现来源：安全自查 / 复测（2026-05）"
            ),
            suggestion=(
                "登录口令传输改用非对称加密或仅依赖 TLS；避免匿名对称密钥下发；渐进下线 AES 口令通道。"
            ),
            proof="接口 PoC：GET /xgszp-boot/sys/getEncryptedString；加固后应以 RSA 链路为主并复测登录。",
            fix_detail=(
                "一次性 RSA 会话密钥 + Redis 私钥；前端 jeecgboot-vue3/src/utils/encryption/loginRsa.ts、store/modules/user.ts；后端 LoginRsaCryptoUtil、SysLoginModel.rsaKeyId。"
            ),
            files=(
                "LoginController.java；LoginRsaCryptoUtil.java；SysLoginModel.java；ShiroConfig.java；CommonConstant.java；"
                "jeecgboot-vue3/src/utils/encryption/loginRsa.ts；jeecgboot-vue3/src/store/modules/user.ts；jeecgboot-vue3/package.json"
            ),
            self_test=(
                "1）GET /sys/loginRsaPublicKey 返回 rsaKeyId、publicKey（PEM）；2）RSA 加密后 POST /sys/login 成功；"
                "3）Redis 中 LOGIN_RSA:PRIV:* 使用后删除；4）旧客户端无 rsaKeyId 仍可登录（若仍启用 AES）。"
            ),
            retest="☑ 建议复测：匿名不再依赖明文 AES 口令链路；登录与会话审计（待安全签字）。",
        ),
    ),
    (
        "信工所招聘管理系统 — 漏洞记录及修复情况记录表（第 39 / 41 条｜已修复）",
        base_rows(
            name="通用上传接口允许 SVG 嵌脚本（存储型 XSS 风险）",
            vuln_id="VULN-039",
            level="中危",
            description=(
                "【修复状态：已修复】POST /sys/common/upload 使用 SsrfFileTypeFilter 白名单，历史包含 svg；"
                "SVG 为 XML，可嵌入 script/on*，同源预览可导致存储型 XSS。与 VULN-006/VULN-021 同一上传面加固。\n"
                "【本次修复】从 FILE_TYPE_WHITE_LIST 移除 svg（上传与共用该表的下载校验一并收紧）。\n"
                "发现来源：安全自查 / 复测（2026-05）"
            ),
            suggestion="禁止用户可控 SVG 上传或服务端消毒；静态矢量资源改为打包下发。",
            proof="构造含 script 的 .svg 经 /sys/common/upload 上传成功（修复前）；修复后应返回非法类型。",
            fix_detail="SsrfFileTypeFilter 去掉 svg；回归 pdf/图片上传。",
            files="jeecg-boot-base-core/.../SsrfFileTypeFilter.java",
            self_test="求职者或管理员上传 .svg 应被拒绝；.pdf/.jpg/.png 仍成功。",
            retest="☑ 建议复测：svg 上传拒绝（待安全签字）。",
        ),
    ),
    (
        "信工所招聘管理系统 — 漏洞记录及修复情况记录表（第 40 / 41 条｜已修复）",
        base_rows(
            name="多数据源写操作接口缺失 Shiro 权限（低权限可添加数据源）",
            vuln_id="VULN-040",
            level="高危",
            description=(
                "【修复状态：已修复】POST /sys/dataSource/add（及 edit/delete 等）缺少 @RequiresPermissions，"
                "仅依赖 JdbcSecurityUtil 校验 JDBC URL，任意持有 JWT 的低权限账号仍可写入数据源配置。\n"
                "【修复】add/edit/delete/deleteBatch/queryById/exportXls/importExcel 补齐 system:datasource:*；"
                "Flyway V3.9.3_0__xgs_security_quartz_datasource_buttons.sql 写入按钮权限并按已有 datasource:list 角色授权。\n"
                "发现来源：安全自查 / 复测（2026-05）"
            ),
            suggestion="所有数据源管理接口必须与菜单权限绑定；求职者角色禁止访问系统管理 API。",
            proof="求职者 Token POST /sys/dataSource/add 返回成功（修复前）；修复后应无权限。",
            fix_detail="SysDataSourceController 注解 + Flyway 权限种子数据。",
            files=(
                "SysDataSourceController.java；"
                "flyway/sql/mysql/V3.9.3_0__xgs_security_quartz_datasource_buttons.sql"
            ),
            self_test="求职者账号调用 dataSource/add 应 403；管理员具备 datasource:list 及新按钮权限后可正常维护。",
            retest="☑ 建议复测：低权限写入数据源不可复现（待安全签字）。",
        ),
    ),
    (
        "信工所招聘管理系统 — 漏洞记录及修复情况记录表（第 41 / 41 条｜已修复）",
        base_rows(
            name="定时任务管理列表等接口未授权访问（信息泄露）",
            vuln_id="VULN-041",
            level="中危",
            description=(
                "【修复状态：已修复】GET /sys/quartzJob/list 及 queryById、exportXls、importExcel 未加 @RequiresPermissions，"
                "任意登录用户可读取调度任务详情。\n"
                "【修复】补齐 system:quartzJob:list、queryById、exportXls、importExcel；"
                "同上 Flyway 脚本向已具备 quartzJob:add 的角色批量授予新权限。\n"
                "发现来源：安全自查 / 复测（2026-05）"
            ),
            suggestion="定时任务所有读写接口纳入 RBAC；默认拒绝求职者。",
            proof="低权限 JWT GET /sys/quartzJob/list 返回分页数据（修复前）。",
            fix_detail="QuartzJobController 注解 + Flyway。",
            files=(
                "QuartzJobController.java；"
                "flyway/sql/mysql/V3.9.3_0__xgs_security_quartz_datasource_buttons.sql"
            ),
            self_test="无 list 权限用户 GET /sys/quartzJob/list 应拒绝；具备权限的管理员分页正常。",
            retest="☑ 建议复测：未授权访问 Quartz 列表不可复现（待安全签字）。",
        ),
    ),
]


def replace_in_paragraphs(doc: Document, old: str, new: str) -> int:
    n = 0
    for p in doc.paragraphs:
        if old in p.text:
            p.text = p.text.replace(old, new)
            n += 1
    return n


def replace_in_tables(doc: Document, old: str, new: str) -> None:
    from docx.oxml.exceptions import InvalidXmlError

    for tbl in doc.tables:
        try:
            for row in tbl.rows:
                for cell in row.cells:
                    for old_p in cell.paragraphs:
                        if old in old_p.text:
                            old_p.text = old_p.text.replace(old, new)
        except InvalidXmlError:
            continue


def find_anchor_paragraph(doc: Document):
    for p in doc.paragraphs:
        if "漏洞记录补充条（报告名" in p.text or "漏洞记录补充条（报告名 岗位发布界面 XSS）" in p.text:
            return p
        if "漏洞记录补充条" in p.text and "岗位发布" in p.text:
            return p
    for p in doc.paragraphs:
        if "【总体修复策略" in p.text and "同步至本汇总表" in p.text:
            return p
    return None


def insert_section_before(paragraph, doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    """在 paragraph 之前插入：标题（居中）、空行、表格（顺序与既有主表条一致）。"""
    p_title = doc.add_paragraph(title)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_el = p_title._element
    title_el.getparent().remove(title_el)

    p_blank = doc.add_paragraph("")
    blank_el = p_blank._element
    blank_el.getparent().remove(blank_el)

    tbl = doc.add_table(rows=len(rows), cols=2)
    try:
        tbl.style = "Table Grid"
    except (KeyError, ValueError):
        pass
    for i, (c0, c1) in enumerate(rows):
        tbl.rows[i].cells[0].text = c0
        tbl.rows[i].cells[1].text = c1
    tbl_el = tbl._tbl
    tbl_el.getparent().remove(tbl_el)

    anchor_el = paragraph._element
    parent = anchor_el.getparent()
    idx = parent.index(anchor_el)
    parent.insert(idx, title_el)
    parent.insert(idx + 1, blank_el)
    parent.insert(idx + 2, tbl_el)


def append_reference_alias_row(doc: Document) -> None:
    """文首对照表（第一张表）追加一行说明。"""
    if not doc.tables:
        return
    t0 = doc.tables[0]
    for row in t0.rows:
        if row.cells and "登录 AES 密钥匿名下发" in row.cells[0].text:
            return
    row = t0.add_row()
    cells = row.cells
    if len(cells) >= 5:
        cells[0].text = "登录 AES 密钥匿名下发 / RSA 口令加固（自查）"
        cells[1].text = "登录口令 AES 密钥匿名下发（getEncryptedString）"
        cells[2].text = "VULN-038"
        cells[3].text = "中危"
        cells[4].text = "与 VULN-004（签名密钥）区分；详见主表 VULN-038。"
    elif len(cells) >= 4:
        cells[0].text = "登录 RSA 口令加固（自查）"
        cells[1].text = "VULN-038～041 见主表新增条"
        cells[2].text = "见主表"
        cells[3].text = "—"


def extend_footer_vuln_list(doc: Document) -> None:
    append = ", VULN-038, VULN-039, VULN-040, VULN-041"
    for p in doc.paragraphs:
        if (
            "【已修复】编号：" in p.text
            and "VULN-001" in p.text
            and "VULN-038" not in p.text
        ):
            p.text = p.text.rstrip() + append
            return


def main() -> int:
    path = DOC_PATH
    if len(sys.argv) > 1:
        path = Path(sys.argv[1])
    if not path.is_file():
        print(f"文件不存在: {path}", file=sys.stderr)
        return 1

    doc = Document(str(path))

    joined = "\n".join(p.text for p in doc.paragraphs)
    already_have_new_sections = "（第 38 / 41 条｜已修复）" in joined

    if not already_have_new_sections:
        old_suffix = "/ 37 条｜已修复）"
        new_suffix = "/ 41 条｜已修复）"
        replace_in_paragraphs(doc, old_suffix, new_suffix)
        replace_in_tables(doc, old_suffix, new_suffix)
        replace_in_paragraphs(doc, "第1/37 条之体系）：已修复 37 条", "第1/41 条之体系）：已修复 41 条")
        replace_in_paragraphs(doc, "合计书面可检索条目 42 条", "合计书面可检索条目 46 条")

        anchor = find_anchor_paragraph(doc)
        if anchor is None:
            print("未找到插入锚点段落（漏洞记录补充条或「总体修复策略」）", file=sys.stderr)
            return 2

        for title, rows in SECTIONS:
            insert_section_before(anchor, doc, title, rows)
    else:
        print(f"主表已含第38～41条，跳过插入与主表条数替换: {path}")

    replace_in_paragraphs(doc, "• 已修复：37 条", "• 已修复：41 条")
    extend_footer_vuln_list(doc)

    doc.save(str(path))
    print(f"已更新: {path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
